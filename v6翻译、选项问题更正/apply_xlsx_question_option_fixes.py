from __future__ import annotations

import csv
import re
import shutil
from copy import copy, deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


ROOT = Path(__file__).resolve().parent
SOURCE_DIR = ROOT / "英文版解析_中文版覆盖"
OUTPUT_DIR = ROOT / "英文版题干选项_按复核意见修订"
REVIEW_PATH = ROOT / "CAMS题库问题清单_复核版.xlsx"


@dataclass
class Question:
    number: int
    question_paragraph: object
    option_paragraphs: dict[str, object]
    answer_paragraph: object
    pre_answer_paragraphs: list


def parse_questions(document: Document) -> list[Question]:
    paragraphs = document.paragraphs
    texts = [paragraph.text.strip() for paragraph in paragraphs]
    answer_indexes = [
        index for index, text in enumerate(texts) if re.match(r"^答案\s*[:：]", text)
    ]
    starts: list[int] = []
    previous_answer = -1
    for answer_index in answer_indexes:
        option_a_indexes = [
            index
            for index in range(previous_answer + 1, answer_index)
            if re.match(r"^A[.．、]", texts[index], re.IGNORECASE)
        ]
        if not option_a_indexes:
            raise ValueError(f"Cannot find option A before answer paragraph {answer_index}")
        question_indexes = [
            index
            for index in range(previous_answer + 1, option_a_indexes[-1])
            if re.match(r"^\d+[.．、]", texts[index])
        ]
        if not question_indexes:
            raise ValueError(f"Cannot find question before answer paragraph {answer_index}")
        starts.append(question_indexes[-1])
        previous_answer = answer_index

    questions: list[Question] = []
    for number, (start, answer_index) in enumerate(zip(starts, answer_indexes), start=1):
        option_paragraphs = {}
        for paragraph in paragraphs[start + 1 : answer_index]:
            match = re.match(r"^([A-Z])[.．、]", paragraph.text.strip(), re.IGNORECASE)
            if match:
                option_paragraphs[match.group(1).upper()] = paragraph
        questions.append(
            Question(
                number=number,
                question_paragraph=paragraphs[start],
                option_paragraphs=option_paragraphs,
                answer_paragraph=paragraphs[answer_index],
                pre_answer_paragraphs=paragraphs[start:answer_index],
            )
        )
    return questions


def set_paragraph_text(paragraph, value: str) -> None:
    first_run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_run_properties = deepcopy(paragraph.runs[0]._r.rPr)
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(value)
    if first_run_properties is not None:
        run._r.insert(0, first_run_properties)


def field_paragraph(question: Question, field: str):
    if field == "question":
        return question.question_paragraph
    if field not in question.option_paragraphs:
        raise KeyError(f"Question {question.number} has no option {field}")
    return question.option_paragraphs[field]


def main() -> None:
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    OUTPUT_DIR.mkdir()

    source_paths = sorted(SOURCE_DIR.glob("*.docx"))
    documents = {}
    questions = {}
    original_counts = {}
    for source_path in source_paths:
        output_path = OUTPUT_DIR / source_path.name
        shutil.copy2(source_path, output_path)
        document = Document(output_path)
        documents[source_path.name] = document
        questions[source_path.name] = parse_questions(document)
        original_counts[source_path.name] = len(questions[source_path.name])

    changes: list[dict[str, str | int]] = []

    def change(
        file_name: str,
        question_number: int,
        field: str,
        new_text: str,
        excel_rows: str,
        reason: str,
    ) -> None:
        question = questions[file_name][question_number - 1]
        paragraph = field_paragraph(question, field)
        before = paragraph.text
        if before == new_text:
            return
        set_paragraph_text(paragraph, new_text)
        changes.append(
            {
                "英文文件": file_name,
                "当前题号": question_number,
                "字段": "题干" if field == "question" else f"选项{field}",
                "Excel行": excel_rows,
                "修改原因": reason,
                "修改前": before,
                "修改后": new_text,
            }
        )

    # Second chapter: terminology and machine-translation corrections.
    change(
        "英文版第二章.docx",
        1,
        "question",
        "1.Which legal or illegal activity, when combined with willful blindness and an international criminal element, may lead to the filing of a suspicious activity report?",
        "7",
        "Use the textbook term 'willful blindness' and restore natural English.",
    )
    change(
        "英文版第二章.docx",
        1,
        "A",
        "A.Illegal activity whose proceeds could result in money laundering charges.",
        "7",
        "Rewrite the machine-translated option without changing its meaning.",
    )
    change(
        "英文版第二章.docx",
        1,
        "B",
        "B.An interface that forms part of a suspicious transaction monitoring system.",
        "7",
        "Rewrite the machine-translated option without changing its meaning.",
    )
    change(
        "英文版第二章.docx",
        1,
        "C",
        "C.A specific illegal activity involving the use of a concentration account to defraud customers who are not directly associated with that account.",
        "7",
        "Rewrite the machine-translated option without changing its meaning.",
    )
    change(
        "英文版第二章.docx",
        50,
        "question",
        '50.What is "nesting" in correspondent banking?',
        "12",
        "Use the standard correspondent-banking term.",
    )
    change(
        "英文版第二章.docx",
        50,
        "A",
        "A.A correspondent bank provides upstream correspondent services to other financial institutions.",
        "12",
        "Replace nonstandard machine-translated terminology.",
    )
    change(
        "英文版第二章.docx",
        50,
        "B",
        "B.A respondent bank provides downstream correspondent services to other financial institutions.",
        "12",
        "Use the textbook definition of nesting.",
    )
    change(
        "英文版第二章.docx",
        50,
        "C",
        "C.The bank shares many customers with other local banks.",
        "12",
        "Correct capitalization and grammar.",
    )
    change(
        "英文版第二章.docx",
        51,
        "question",
        "51.How does a payable-through account (PTA) expose a correspondent bank to money laundering risk?",
        "13, 14",
        "Replace 'remittance account' with the textbook term PTA.",
    )
    change(
        "英文版第二章.docx",
        51,
        "A",
        "A.It may be used by a respondent bank that has not established adequate due diligence requirements.",
        "13, 14",
        "Replace 'entrusting bank' with 'respondent bank'.",
    )
    change(
        "英文版第二章.docx",
        51,
        "B",
        "B.It uses checks with numerical codes to identify the relevant subaccounts and their owners.",
        "13",
        "Correct grammar while preserving the option's meaning.",
    )
    change(
        "英文版第二章.docx",
        51,
        "C",
        "C.Subaccount holders may have direct access to funds in the account.",
        "13",
        "Express the core PTA risk accurately.",
    )
    change(
        "英文版第二章.docx",
        51,
        "D",
        "D.It may be used by subaccount holders that are multinational financial institutions.",
        "13",
        "Correct grammar while preserving the option's meaning.",
    )
    change(
        "英文版第二章.docx",
        52,
        "question",
        "52.Which policy is necessary for payable-through account services to maintain a sound AML/CFT framework?",
        "4, 15",
        "Remove the invalid transliteration 'Huitong'.",
    )
    change(
        "英文版第二章.docx",
        52,
        "A",
        "A.Provide employees with practical guidance on identifying the ultimate beneficial owners of the respondent bank.",
        "4, 15",
        "Use standard correspondent-banking terminology.",
    )
    change(
        "英文版第二章.docx",
        52,
        "B",
        "B.Exercise comprehensive oversight of respondent banks and their AML/CFT frameworks.",
        "4, 15",
        "Use standard correspondent-banking terminology.",
    )
    change(
        "英文版第二章.docx",
        52,
        "C",
        "C.Authorize the correspondent bank to terminate relationships with high-risk payable-through account customers.",
        "4, 15",
        "Remove machine-translated terminology and repair the sentence.",
    )
    change(
        "英文版第二章.docx",
        52,
        "D",
        "D.Require the correspondent bank to notify the respondent bank's compliance officer of red flags before taking action.",
        "4, 15",
        "Remove the invalid transliteration and repair the sentence.",
    )
    change(
        "英文版第二章.docx",
        53,
        "question",
        "53.Which circumstance makes a payable-through account most vulnerable to use as a money laundering channel?",
        "13",
        "Replace 'remittance account' with the textbook term PTA.",
    )
    change(
        "英文版第二章.docx",
        53,
        "A",
        "A.It is maintained for a foreign financial institution, such as a bank.",
        "13",
        "Correct grammar and terminology.",
    )
    change(
        "英文版第二章.docx",
        53,
        "B",
        "B.It does not provide services directly to third parties.",
        "13",
        "Correct grammar and terminology.",
    )
    change(
        "英文版第二章.docx",
        53,
        "C",
        "C.It is maintained for a foreign bank with no physical presence in any country.",
        "13",
        "Use the standard definition of a foreign shell bank.",
    )
    change(
        "英文版第二章.docx",
        53,
        "D",
        "D.It is maintained for a publicly listed foreign private bank authorized to act as an intermediary.",
        "13",
        "Correct grammar and terminology.",
    )
    change(
        "英文版第二章.docx",
        57,
        "question",
        "57.(多选题)A company service provider in Country A established a corporate structure for a customer from Country B, which is known for corruption. The structure includes a holding company in Country A with an account at an international bank. At onboarding, the customer's wealth was estimated at $7.52 million. Shortly thereafter, the customer's father became the president of Country B. Two years later, a routine customer review found that the customer's wealth had grown to $510 million. Which two red flags indicate possible money laundering or terrorist financing? (Choose two.)",
        "10, 24",
        "Correct onboarding language, country-risk wording, currency format, and 'terrorist financing'.",
    )
    change(
        "英文版第二章.docx",
        63,
        "B",
        "B.Check deposits followed by ATM withdrawals using a debit card at a retail store.",
        "11",
        "Correct 'debt card' to 'debit card'.",
    )
    change(
        "英文版第二章.docx",
        66,
        "question",
        "66.What action should a money services business (MSB) take to avoid violating the Bank Secrecy Act (BSA)?",
        "16",
        "Use standard MSB and BSA terminology.",
    )
    change(
        "英文版第二章.docx",
        66,
        "A",
        "A.Ensure that the BSA/AML department directly informs every customer and adopt policies that maximize profitability and customer anonymity.",
        "16",
        "Replace machine-translated AML terminology.",
    )
    change(
        "英文版第二章.docx",
        66,
        "B",
        "B.Develop and maintain an AML program appropriate to the MSB's location, size, nature, and transaction volume.",
        "16",
        "Use standard MSB, BSA, and AML program terminology.",
    )
    change(
        "英文版第二章.docx",
        66,
        "C",
        "C.Allow customers to control what information they provide to the MSB and how long the information is retained.",
        "16",
        "Replace machine-translated MSB terminology.",
    )
    change(
        "英文版第二章.docx",
        66,
        "D",
        "D.Develop and maintain an AML program based solely on the AML staff and financial resources available to the MSB.",
        "16",
        "Replace machine-translated MSB terminology.",
    )
    change(
        "英文版第二章.docx",
        69,
        "question",
        "69.Which aspect of a money services business (MSB) increases its financial crime risk?",
        "16",
        "Use the standard term 'money services business (MSB)'.",
    )
    change(
        "英文版第二章.docx",
        69,
        "D",
        "D.Serving areas with limited access to banking services.",
        "16",
        "Correct grammar while preserving the option's meaning.",
    )
    change(
        "英文版第二章.docx",
        43,
        "question",
        questions["英文版第二章.docx"][42].question_paragraph.text.replace("wasapproved", "was approved"),
        "31",
        "Insert the missing space.",
    )

    # Third chapter.
    change(
        "英文版第三章.docx",
        24,
        "B",
        "B.Persons residing in, or having funds from, countries with inadequate AML standards.",
        "18",
        "Repair the malformed option.",
    )
    change(
        "英文版第三章.docx",
        146,
        "question",
        "146.What does designating a country as a jurisdiction of primary money laundering concern allow the U.S. government to do?",
        "21",
        "Correct 'designing' and use the statutory term.",
    )
    change(
        "英文版第三章.docx",
        146,
        "D",
        "D.Ensure that the country is included in sanctions programs administered by the Office of Foreign Assets Control.",
        "21",
        "Correct the OFAC name and capitalization.",
    )
    change(
        "英文版第三章.docx",
        175,
        "question",
        "175.Which statement most accurately describes who must comply with Office of Foreign Assets Control (OFAC) sanctions?",
        "22",
        "Correct OFAC capitalization and grammar.",
    )
    change(
        "英文版第三章.docx",
        175,
        "D",
        "D.All U.S. citizens and permanent residents while located in the U.S., all U.S.-incorporated entities and their foreign branches, and all individuals within the United States.",
        "22",
        "Remove source text accidentally appended to the option.",
    )

    # Fourth, fifth, and sixth chapters.
    change(
        "英文版第四章.docx",
        75,
        "A",
        "A.Not involved in the organization's AML compliance program and having a reporting line to the board of directors or a committee thereof.",
        "23",
        "Correct capitalization, possessive form, and wording.",
    )
    change(
        "英文版第四章.docx",
        152,
        "question",
        "152.A quarterly review is conducted on a retail customer's account at a bank in a jurisdiction with currency reporting thresholds. Several large deposits of financial instruments drawn on other institutions, each below the reporting threshold, are identified. This activity is inconsistent with the account's historical profile. A suspicious transaction report is most likely to be filed if which additional activity occurred?",
        "32",
        "Repair OCR, possessives, and sentence structure.",
    )
    change(
        "英文版第四章.docx",
        188,
        "question",
        questions["英文版第四章.docx"][187]
        .question_paragraph.text.replace("totaling$$", "totaling $")
        .replace("$$", "$")
        .replace("$50,000.Account", "$50,000. Account")
        .replace("received.What", "received. What"),
        "33",
        "Remove duplicate dollar signs and restore surrounding spaces.",
    )
    change(
        "英文版第五章.docx",
        86,
        "question",
        questions["英文版第五章.docx"][85].question_paragraph.text.replace("FinanciaI", "Financial"),
        "25",
        "Correct capital-I OCR in 'Financial'.",
    )
    change(
        "英文版第五章.docx",
        108,
        "E",
        "E.Examine financial institutions to ensure that they comply with anti-money laundering regulations.",
        "26",
        "Remove the erroneous word 'alcohol' and repair the option.",
    )
    change(
        "英文版第五章.docx",
        89,
        "question",
        "89.An AML specialist at a small bank has identified suspicious activity at a branch located in an area known for drug trafficking. The investigation reveals that the suspicious transactions occurred during the past three months and were processed by the same teller. The teller did not file an internal unusual activity report. A review of personnel files shows that the teller has been a trusted employee for more than 15 years, has an impeccable work record, and has completed several AML training sessions. The specialist recently learned that the employee's daughter has a rare disease and is undergoing very expensive treatment. What should the specialist recommend regarding the teller's failure to report the unusual activity?",
        "28",
        "Repair concatenation, punctuation, spelling, and sentence structure.",
    )
    change(
        "英文版第六章.docx",
        9,
        "question",
        questions["英文版第六章.docx"][8].question_paragraph.text.replace("theMLRO", "the MLRO"),
        "29",
        "Insert the missing space.",
    )

    # Current-state global cleanup requested by the review workbook.
    for file_name, file_questions in questions.items():
        for question in file_questions:
            for paragraph in question.pre_answer_paragraphs:
                before = paragraph.text
                after = before.replace("$$", "$").replace("Fls", "FIs")
                if paragraph is question.question_paragraph:
                    after = after.replace("(多选题)(多选题)", "(多选题)")
                if after != before:
                    set_paragraph_text(paragraph, after)
                    reasons = []
                    excel_rows = []
                    if "$$" in before:
                        reasons.append("Replace duplicate dollar signs")
                        excel_rows.append("33")
                    if "Fls" in before:
                        reasons.append("Correct lowercase-l/capital-I OCR to FIs")
                        excel_rows.append("20")
                    if "(多选题)(多选题)" in before:
                        reasons.append("Remove the duplicate multiple-choice marker")
                        excel_rows.append("34")
                    field_match = re.match(r"^([A-Z])[.．、]", after.strip())
                    changes.append(
                        {
                            "英文文件": file_name,
                            "当前题号": question.number,
                            "字段": f"选项{field_match.group(1)}" if field_match else "题干",
                            "Excel行": ", ".join(excel_rows),
                            "修改原因": "; ".join(reasons),
                            "修改前": before,
                            "修改后": after,
                        }
                    )

    # The OCR token in Q54 appears in a singular possessive phrase, not as the
    # plural abbreviation used elsewhere.
    q54 = questions["英文版第三章.docx"][53]
    q54_before = q54.question_paragraph.text
    q54_after = q54_before.replace("an FIs AML controls", "an FI's AML controls")
    if q54_after != q54_before:
        set_paragraph_text(q54.question_paragraph, q54_after)
        changes.append(
            {
                "英文文件": "英文版第三章.docx",
                "当前题号": 54,
                "字段": "题干",
                "Excel行": "20",
                "修改原因": "Use the singular possessive form required by the sentence.",
                "修改前": q54_before,
                "修改后": q54_after,
            }
        )

    # The old two single-answer/multiple-choice-marker records no longer exist;
    # repair the one equivalent issue found in the current document set.
    current_q85 = questions["英文版第三章.docx"][84]
    q85_before = current_q85.question_paragraph.text
    q85_after = (
        "85.Under the European Union's Fourth Anti-Money Laundering Directive, "
        "what cash transaction threshold applies to natural or legal persons trading in goods?"
    )
    if q85_after != q85_before:
        set_paragraph_text(current_q85.question_paragraph, q85_after)
        changes.append(
            {
                "英文文件": "英文版第三章.docx",
                "当前题号": 85,
                "字段": "题干",
                "Excel行": "8, 9（当前状态同类项）",
                "修改原因": "The question asks for one threshold and the answer contains one option.",
                "修改前": q85_before,
                "修改后": q85_after,
            }
        )

    for file_name, document in documents.items():
        document.save(OUTPUT_DIR / file_name)

    with (OUTPUT_DIR / "题干选项修订记录.csv").open(
        "w", encoding="utf-8-sig", newline=""
    ) as handle:
        writer = csv.DictWriter(handle, fieldnames=list(changes[0]))
        writer.writeheader()
        writer.writerows(changes)

    # Add current-location and processing columns to a copy of the review workbook.
    workbook = load_workbook(REVIEW_PATH)
    worksheet = workbook.active
    extra_headers = ["current_location", "processing_status", "processing_note"]
    first_extra_column = worksheet.max_column + 1
    for offset, header in enumerate(extra_headers):
        cell = worksheet.cell(1, first_extra_column + offset, header)
        source = worksheet.cell(1, first_extra_column - 1)
        if source.has_style:
            cell._style = copy(source._style)

    results = {
        2: ("", "本轮未处理", "仅涉及答案/解析一致性"),
        3: ("", "待人工确认", "涉及答案变更，需核对原题"),
        4: ("英文第二章 Q52", "部分完成", "术语已修订；答案A/B争议仍待确认"),
        5: ("", "本轮未处理", "仅涉及答案/解析一致性"),
        6: ("", "本轮未处理", "涉及答案及解析重写"),
        7: ("英文第二章 Q1", "已完成", "按当前Word定位；保留原三选项结构"),
        8: ("当前同类项：英文第三章 Q85", "已完成", "旧题已不存在；修正当前同类题"),
        9: ("当前同类项：英文第三章 Q85", "已完成", "旧题已不存在；修正当前同类题"),
        10: ("英文第二章 Q57", "已完成", "当前仅保留一题，已统一术语和表达"),
        11: ("英文第二章 Q63", "已完成", "debt card -> debit card"),
        12: ("英文第二章 Q50", "已完成", "按当前Word定位"),
        13: ("英文第二章 Q51、Q53", "已完成", "统一为 payable-through account (PTA)"),
        14: ("英文第二章 Q51", "已完成", "entrusting bank -> respondent bank"),
        15: ("英文第二章 Q52", "已完成", "删除 Huitong 音译"),
        16: ("英文第二章 Q66、Q69", "已完成", "统一 MSB、BSA、AML program"),
        17: ("英文第二章 Q72", "无需修改英文", "当前英文 surrender charges 正确；意见针对中文回译"),
        18: ("英文第三章 Q24", "已完成", "按当前Word定位"),
        19: ("英文第三章 Q181", "待人工确认", "A/D选项区分度涉及原题内容"),
        20: ("英文第三章 Q54、Q141、Q142", "已完成", "修复当前所有 Fls OCR"),
        21: ("英文第三章 Q146", "已完成", "按当前Word定位"),
        22: ("英文第三章 Q175", "已完成", "修复OFAC名称并删除来源文字"),
        23: ("英文第四章 Q75", "已完成", "按当前Word定位"),
        24: ("英文第二章 Q57", "部分完成", "英文金额格式已统一；旧版重复解析已不存在"),
        25: ("英文第五章 Q86", "已完成", "FinanciaI -> Financial"),
        26: ("英文第五章 Q108", "已完成", "删除 alcohol OCR错误"),
        27: ("英文第五章 Q107", "当前已正确", "已使用 SAR supporting documentation；重复题已不存在"),
        28: ("英文第五章 Q89", "已完成", "重写拼接、拼写错误题干"),
        29: ("英文第六章 Q9", "已完成", "theMLRO -> the MLRO"),
        30: ("英文第六章 Q3", "待人工确认", "涉及答案集合和取证依据"),
        31: ("英文第二章 Q43", "已完成", "wasapproved -> was approved"),
        32: ("英文第四章 Q152", "已完成", "修复OCR、所有格及句子结构"),
        33: ("当前6题", "已完成", "当前Word共发现并修复6处 $$"),
        34: ("英文第二章 Q168；英文第三章 Q197、Q198", "已完成", "当前Word共发现并修复3处重复标记"),
        35: ("多组", "待人工确认", "本轮不删除题目"),
    }
    for row_number in range(2, worksheet.max_row + 1):
        location, status, note = results[row_number]
        worksheet.cell(row_number, first_extra_column, location)
        worksheet.cell(row_number, first_extra_column + 1, status)
        worksheet.cell(row_number, first_extra_column + 2, note)
    worksheet.auto_filter.ref = worksheet.dimensions
    workbook.save(OUTPUT_DIR / "CAMS题库问题清单_处理结果.xlsx")

    # Build a compact, question-level checklist for import and manual review.
    field_changes: dict[tuple[str, int, str], dict[str, str]] = {}
    for item in changes:
        key = (str(item["英文文件"]), int(item["当前题号"]), str(item["字段"]))
        if key not in field_changes:
            field_changes[key] = {
                "before": str(item["修改前"]),
                "after": str(item["修改后"]),
                "excel_rows": str(item["Excel行"]),
            }
        else:
            field_changes[key]["after"] = str(item["修改后"])
            existing_rows = field_changes[key]["excel_rows"].split(", ")
            for row in str(item["Excel行"]).split(", "):
                if row not in existing_rows:
                    existing_rows.append(row)
            field_changes[key]["excel_rows"] = ", ".join(existing_rows)

    question_changes: dict[tuple[str, int], list[tuple[str, dict[str, str]]]] = {}
    for (file_name, question_number, field), detail in field_changes.items():
        question_changes.setdefault((file_name, question_number), []).append((field, detail))

    chapter_order = {
        "英文版第二章.docx": 2,
        "英文版第三章.docx": 3,
        "英文版第四章.docx": 4,
        "英文版第五章.docx": 5,
        "英文版第六章.docx": 6,
    }
    summary_workbook = Workbook()
    summary_sheet = summary_workbook.active
    summary_sheet.title = "题目级修改清单"
    summary_headers = [
        "序号",
        "修订版DOCX",
        "当前题号",
        "修改字段",
        "Excel意见行",
        "修改前",
        "修改后",
        "后台处理状态",
        "核对备注",
    ]
    summary_sheet.append(summary_headers)
    sorted_questions = sorted(
        question_changes.items(), key=lambda item: (chapter_order[item[0][0]], item[0][1])
    )
    for sequence, ((file_name, question_number), items) in enumerate(sorted_questions, start=1):
        field_order = {"题干": 0, "选项A": 1, "选项B": 2, "选项C": 3, "选项D": 4, "选项E": 5, "选项F": 6}
        items.sort(key=lambda item: field_order.get(item[0], 99))
        fields = "、".join(field for field, _ in items)
        excel_rows: list[str] = []
        before_lines = []
        after_lines = []
        for field, detail in items:
            before_lines.append(f"[{field}] {detail['before']}")
            after_lines.append(f"[{field}] {detail['after']}")
            for row in detail["excel_rows"].split(", "):
                if row and row not in excel_rows:
                    excel_rows.append(row)
        summary_sheet.append(
            [
                sequence,
                file_name,
                question_number,
                fields,
                "、".join(excel_rows),
                "\n".join(before_lines),
                "\n".join(after_lines),
                "待修改",
                "",
            ]
        )

    header_fill = PatternFill("solid", fgColor="1F4E78")
    for cell in summary_sheet[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row in summary_sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    widths = [8, 22, 10, 24, 18, 85, 85, 14, 30]
    for column, width in enumerate(widths, start=1):
        summary_sheet.column_dimensions[summary_sheet.cell(1, column).column_letter].width = width
    summary_sheet.freeze_panes = "A2"
    summary_sheet.auto_filter.ref = summary_sheet.dimensions
    summary_workbook.save(OUTPUT_DIR / "题干选项修改清单.xlsx")

    print(f"documents={len(documents)} changes={len(changes)}")
    print(OUTPUT_DIR)


if __name__ == "__main__":
    main()
