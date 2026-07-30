from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CAMS题库项目说明与运营使用指南.md"
OUTPUT_DIR = ROOT / "题库解析SOP_DOCX"
OUTPUT = OUTPUT_DIR / "CAMS题库项目说明与运营使用指南.docx"
ASSET_DIR = OUTPUT_DIR / "_assets"
DIAGRAM = ASSET_DIR / "CAMS题库项目说明与运营使用指南_主流程图.png"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
BLUE = "1F4E79"
BLACK = "000000"


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str = BLACK) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before: float = 0, after: float = 5, line: float = 1.38) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def parse_inline(paragraph, text: str, size: float = 10.5, color: str = BLACK) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color=BLUE)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(max(size - 0.5, 8))
            run.font.color.rgb = RGBColor.from_string(BLACK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def make_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if FONT_PATH.exists():
        return ImageFont.truetype(str(FONT_PATH), size=size, index=0)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        current = ""
        for char in paragraph:
            candidate = current + char
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = char
        if current:
            lines.append(current)
    return lines or [""]


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str) -> None:
    draw.rounded_rectangle(box, radius=8, fill="#FFFFFF", outline="#111111", width=4)
    x1, y1, x2, y2 = box
    title_font = make_font(30, bold=True)
    body_font = make_font(24)
    draw.text((x1 + 28, y1 + 24), title, font=title_font, fill="#1F4E79")
    y = y1 + 76
    for line in wrap_text(draw, body, body_font, x2 - x1 - 56):
        draw.text((x1 + 28, y), line, font=body_font, fill="#000000")
        y += 34


def draw_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]]) -> None:
    draw.line(points, fill="#111111", width=5, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        draw.polygon([(x2, y2), (x2 - 18 * direction, y2 - 11), (x2 - 18 * direction, y2 + 11)], fill="#111111")
    else:
        direction = 1 if y2 > y1 else -1
        draw.polygon([(x2, y2), (x2 - 11, y2 - 18 * direction), (x2 + 11, y2 - 18 * direction)], fill="#111111")


def create_main_flow(path: Path) -> None:
    width, height = 2400, 1360
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((90, 52), "两种任务，最后都要经过同样核验", font=make_font(46, bold=True), fill="#000000")
    draw.text((90, 116), "新增题复用已有教材内容；完整建库先整理教材。开始核对依据后，两者遵守同一流程。", font=make_font(24), fill="#000000")

    entry_a = (90, 230, 690, 420)
    entry_b = (90, 510, 690, 700)
    evidence = (870, 370, 1470, 560)
    deepseek = (1650, 370, 2250, 560)
    codex = (1650, 790, 2250, 980)
    human = (870, 790, 1470, 980)
    delivery = (90, 790, 690, 980)

    draw_box(draw, entry_a, "入口 A｜新教材、新题库", "整理教材与批量题目\n建立教材原文和知识联系")
    draw_box(draw, entry_b, "入口 B｜现有题库新增题", "接收题源并清洗查重\n复用已经整理好的教材内容")
    draw_box(draw, evidence, "寻找可能有关的教材原文", "结合题干、各选项和相邻知识\n准备可返回原页核对的材料")
    draw_box(draw, deepseek, "AI 初稿（DeepSeek）", "独立判断选项并形成初稿\n结果不直接发布")
    draw_box(draw, codex, "二次核验与重写（Codex）", "重新核题、核对原文和补充依据\n形成提交人工审核的解析")
    draw_box(draw, human, "人工审核", "确认题面、答案、证据和表达\n批准、退回或暂缓")
    draw_box(draw, delivery, "交付与留档", "导出并核对实际展示\n保留来源、版本和审核记录")

    draw_arrow(draw, [(690, 325), (770, 325), (770, 465), (870, 465)])
    draw_arrow(draw, [(690, 605), (770, 605), (770, 465), (870, 465)])
    draw_arrow(draw, [(1470, 465), (1650, 465)])
    draw_arrow(draw, [(1950, 560), (1950, 790)])
    draw_arrow(draw, [(1650, 885), (1470, 885)])
    draw_arrow(draw, [(870, 885), (690, 885)])

    draw.text((90, 1100), "共同原则", font=make_font(29, bold=True), fill="#1F4E79")
    draw.text((90, 1150), "相关原文仍需核对｜AI 初稿不等于定稿｜依据不足时退回或暂缓｜人工审核后才能交付", font=make_font(25), fill="#000000")
    draw.text((90, 1280), "CAMS 题库项目说明与运营使用指南 · 主流程图", font=make_font(21), fill="#000000")
    image.save(path)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    p = doc.add_paragraph()
    set_spacing(p, before=5, after=5)
    run = p.add_run("CAMS 题库项目说明")
    set_run_font(run, size=21, bold=True)

    p = doc.add_paragraph()
    set_spacing(p, after=12)
    run = p.add_run("使用指南｜项目理解、考生价值和新增题处理逻辑")
    set_run_font(run, size=10.5)

    p = doc.add_paragraph()
    set_spacing(p, after=14)
    run = p.add_run("适用说明：本文件不要求读者掌握程序知识。完整建库方法和新增题内部技术执行分别由对应 SOP 承载。")
    set_run_font(run, size=9.5)

    p = doc.add_paragraph()
    set_spacing(p, after=14)
    run = p.add_run("阅读建议：")
    set_run_font(run, size=9.5, bold=True, color=BLUE)
    run = p.add_run("先阅读第一章速览，后续章节按需查阅具体理由、运营表达、备考建议和新增题处理逻辑。")
    set_run_font(run, size=9.5)


def add_callout(doc: Document, lines: list[str]) -> None:
    text = "\n".join(line for line in lines if line.strip())
    if not text:
        return
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.6)
    set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    p = cell.paragraphs[0]
    set_spacing(p, after=0)
    parse_inline(p, text, size=9.8)
    doc.add_paragraph()


def table_ratios(headers: list[str]) -> list[float]:
    count = len(headers)
    if count == 2:
        return [0.38, 0.62]
    if count == 3:
        return [0.28, 0.40, 0.32]
    if count == 4:
        return [0.20, 0.31, 0.29, 0.20]
    return [1 / count] * count


def add_table(doc: Document, lines: list[str]) -> None:
    separator = re.compile(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?")
    rows = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in lines
        if not separator.fullmatch(line.strip())
    ]
    if not rows:
        return
    columns = max(len(row) for row in rows)
    ratios = table_ratios(rows[0])
    total_cm = 16.6
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    table.autofit = False
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(tbl_layout)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.width = Cm(total_cm * ratios[index])
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.18)
            value = values[index] if index < len(values) else ""
            parse_inline(p, value, size=8.1 if columns == 4 else 8.7)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(BLACK)
        if row_index == 0:
            set_repeat_header(row)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=100, start=150, bottom=100, end=150)
    p = cell.paragraphs[0]
    set_spacing(p, after=0, line=1.15)
    for index, line in enumerate(lines):
        run = p.add_run(line + ("\n" if index < len(lines) - 1 else ""))
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.6)
        run.font.color.rgb = RGBColor.from_string(BLACK)
    doc.add_paragraph()


def add_manual_number(doc: Document, number: str, text: str) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=3)
    p.paragraph_format.left_indent = Cm(0.62)
    p.paragraph_format.first_line_indent = Cm(-0.50)
    run = p.add_run(f"{number}. ")
    set_run_font(run)
    parse_inline(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=3)
    parse_inline(p, text)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)

    for name, size in (("Title", 23), ("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.2)):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)


def build_docx() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_main_flow(DIAGRAM)

    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    source_title_skipped = False
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped == "<!-- PAGE BREAK -->":
            doc.add_page_break()
            index += 1
            continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            continue

        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip("> "))
                index += 1
            add_callout(doc, quote_lines)
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            run = p.add_run()
            run.add_picture(str(DIAGRAM), width=Inches(6.35))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(caption, after=8)
            run = caption.add_run(image_match.group(1))
            set_run_font(run, size=8.5)
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and not source_title_skipped:
                source_title_skipped = True
                index += 1
                continue
            style = {1: "Title", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[level]
            size = {1: 20, 2: 15, 3: 12.5, 4: 11.2}[level]
            p = doc.add_paragraph(style=style)
            set_spacing(p, before=12 if level <= 2 else 8, after=6)
            p.paragraph_format.keep_with_next = True
            parse_inline(p, text, size=size)
            index += 1
            continue

        if not stripped or stripped == "---":
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if numbered:
            add_manual_number(doc, numbered.group(1), numbered.group(2))
        elif bullet:
            add_bullet(doc, bullet.group(1))
        else:
            p = doc.add_paragraph()
            set_spacing(p, after=6)
            parse_inline(p, stripped)
        index += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
