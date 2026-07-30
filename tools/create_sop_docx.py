from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "题库解析SOP_DOCX"
ASSET_DIR = OUTPUT_DIR / "_assets"
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")

DOCS = [
    {
        "source": ROOT / "新教材新题库解析撰写SOP.md",
        "output": OUTPUT_DIR / "新教材新题库解析撰写SOP.docx",
        "title": "新教材、新题库解析撰写 SOP",
        "subtitle": "从教材底座到整套题库解析的建设流程",
        "note": "适用说明：面向负责完整建库、解析生产与教研审核的项目成员，说明各阶段的目的、产物和质量门槛。",
        "lanes": [
            ("材料与系统处理", ["确认教材与题源", "整理教材与批量题目", "建立教材证据与知识联系", "寻找依据并形成初稿"]),
            ("Codex 核验与解析", ["重新核题、核证与补证", "重写面向考生的解析"]),
            ("人工审核与交付", ["责任编辑确认", "独立审核人批准或退回", "导出、上线核对与留档"]),
        ],
    },
    {
        "source": ROOT / "新题加入处理SOP.md",
        "output": OUTPUT_DIR / "新题加入处理SOP.docx",
        "title": "现有题库新增题目处理 SOP",
        "subtitle": "复用现有教材和知识底座的增量解析流程",
        "note": "适用说明：面向新教研、内容处理人员和项目接手人，说明新增题怎样核验、审核和发布；不要求掌握程序知识。",
        "lanes": [
            ("材料与系统处理", ["接收新题并保留题源", "清洗、查重与编号", "初检索并形成候选证据"]),
            ("Codex 核验与解析", ["独立核题、核证与补证", "形成可审核解析或退回"]),
            ("人工审核与交付", ["人工审核与分流", "增量导出、上线核对与留档"]),
        ],
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 5, line: float = 1.38) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def parse_inline(paragraph, text: str, size: float = 10.5, color: str = "000000") -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, color="1F4E79")
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(size - 0.5)
            run.font.color.rgb = RGBColor.from_string("000000")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def make_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if FONT_PATH.exists():
        return ImageFont.truetype(str(FONT_PATH), size=size, index=0)
    return ImageFont.load_default()


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines or [text]


def draw_arrow(draw: ImageDraw.ImageDraw, x1: int, y: int, x2: int) -> None:
    color = "#6380A1"
    draw.line((x1, y, x2 - 20, y), fill=color, width=5)
    draw.polygon([(x2 - 20, y - 11), (x2, y), (x2 - 20, y + 11)], fill=color)


def draw_arrow_left(draw: ImageDraw.ImageDraw, x1: int, y: int, x2: int) -> None:
    color = "#6380A1"
    draw.line((x1, y, x2 + 20, y), fill=color, width=5)
    draw.polygon([(x2 + 20, y - 11), (x2, y), (x2 + 20, y + 11)], fill=color)


def create_flow_diagram(path: Path, title: str, steps: list[str]) -> None:
    width, height = 2400, 1280
    image = Image.new("RGB", (width, height), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    title_font = make_font(48, bold=True)
    step_font = make_font(31)
    small_font = make_font(24)
    draw.text((100, 55), title, font=title_font, fill="#183B5B")
    draw.text((100, 125), "解析不是一次生成，而是经过材料整理、证据核验和审核的完整过程。", font=small_font, fill="#58708A")

    cols = 3
    card_w, card_h = 640, 190
    x_positions = [115, 880, 1645]
    y_positions = [245, 555, 865]
    for index, step in enumerate(steps):
        row, logical_col = divmod(index, cols)
        col = logical_col if row % 2 == 0 else cols - 1 - logical_col
        x, y = x_positions[col], y_positions[row]
        fill = "#E9F2F8" if index % 2 == 0 else "#EEF5EC"
        outline = "#5E88AD" if index % 2 == 0 else "#6D9665"
        draw.rounded_rectangle((x, y, x + card_w, y + card_h), radius=20, fill=fill, outline=outline, width=4)
        draw.ellipse((x + 24, y + 26, x + 88, y + 90), fill=outline)
        draw.text((x + 45, y + 34), str(index + 1), font=small_font, fill="white", anchor="ma")
        lines = wrapped_lines(draw, step, step_font, card_w - 135)
        total_h = len(lines) * 42
        text_y = y + (card_h - total_h) // 2
        for line in lines:
            draw.text((x + 110, text_y), line, font=step_font, fill="#183B5B")
            text_y += 42

        if index + 1 < len(steps):
            next_row, next_logical_col = divmod(index + 1, cols)
            next_col = next_logical_col if next_row % 2 == 0 else cols - 1 - next_logical_col
            next_x, next_y = x_positions[next_col], y_positions[next_row]
            if next_row == row:
                if next_x > x:
                    draw_arrow(draw, x + card_w + 22, y + card_h // 2, next_x - 22)
                else:
                    draw_arrow_left(draw, x - 22, y + card_h // 2, next_x + card_w + 22)
            else:
                center_x = x + card_w // 2
                draw.line((center_x, y + card_h + 18, center_x, next_y - 28), fill="#6380A1", width=5)
                draw.polygon([(center_x - 11, next_y - 28), (center_x, next_y - 8), (center_x + 11, next_y - 28)], fill="#6380A1")

    footer_font = make_font(22)
    draw.text((100, height - 58), "CAMS 题库解析 SOP · 流程示意", font=footer_font, fill="#6B7C8D")
    image.save(path)


def draw_line_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]]) -> None:
    color = "#4C6A82"
    draw.line(points, fill=color, width=4, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        draw.polygon([(x2, y2), (x2 - 15 * direction, y2 - 9), (x2 - 15 * direction, y2 + 9)], fill=color)
    else:
        direction = 1 if y2 >= y1 else -1
        draw.polygon([(x2, y2), (x2 - 9, y2 - 15 * direction), (x2 + 9, y2 - 15 * direction)], fill=color)


def create_swimlane_diagram(path: Path, title: str, lanes: list[tuple[str, list[str]]]) -> None:
    width, height = 2400, 1260
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = make_font(44, bold=True)
    lane_font = make_font(31, bold=True)
    step_font = make_font(29)
    small_font = make_font(22)
    draw.text((92, 48), title, font=title_font, fill="#173A59")
    draw.text((92, 112), "系统准备材料，Codex 核验证据，人工决定是否交付。", font=small_font, fill="#5E6B75")

    lane_y, lane_h = 210, 915
    lane_gap, lane_x, lane_w = 42, 82, 718
    step_num = 1
    lane_boxes: list[list[tuple[int, int, int, int]]] = []
    for lane_index, (lane_title, steps) in enumerate(lanes):
        x = lane_x + lane_index * (lane_w + lane_gap)
        draw.rounded_rectangle((x, lane_y, x + lane_w, lane_y + lane_h), radius=12, fill="#F7F8FA", outline="#AAB7C2", width=3)
        draw.rectangle((x, lane_y, x + lane_w, lane_y + 78), fill="#E8EDF1")
        draw.text((x + 28, lane_y + 21), lane_title, font=lane_font, fill="#173A59")
        top = lane_y + 125
        box_h = 145
        gap = 55 if len(steps) <= 3 else 34
        boxes: list[tuple[int, int, int, int]] = []
        for step in steps:
            box = (x + 36, top, x + lane_w - 36, top + box_h)
            boxes.append(box)
            draw.rounded_rectangle(box, radius=10, fill="#FFFFFF", outline="#6F8495", width=3)
            draw.ellipse((box[0] + 18, box[1] + 36, box[0] + 72, box[1] + 90), fill="#4C6A82")
            draw.text((box[0] + 45, box[1] + 63), str(step_num), font=small_font, fill="#FFFFFF", anchor="mm")
            lines = wrapped_lines(draw, step, step_font, box[2] - box[0] - 112)
            text_y = box[1] + (box_h - len(lines) * 39) // 2
            for line in lines:
                draw.text((box[0] + 94, text_y), line, font=step_font, fill="#173A59")
                text_y += 39
            top += box_h + gap
            step_num += 1
        for prior, following in zip(boxes, boxes[1:]):
            center_x = (prior[0] + prior[2]) // 2
            draw_line_arrow(draw, [(center_x, prior[3] + 10), (center_x, following[1] - 12)])
        lane_boxes.append(boxes)

    for current, following in zip(lane_boxes, lane_boxes[1:]):
        start, end = current[-1], following[0]
        start_pt = (start[2] + 10, (start[1] + start[3]) // 2)
        end_pt = (end[0] - 12, (end[1] + end[3]) // 2)
        middle_x = (start_pt[0] + end_pt[0]) // 2
        draw_line_arrow(draw, [start_pt, (middle_x, start_pt[1]), (middle_x, end_pt[1]), end_pt])

    draw.text((92, height - 72), "CAMS 题库解析 SOP · 阶段泳道图", font=small_font, fill="#6B7C8D")
    image.save(path)


def add_cover(doc: Document, title: str, subtitle: str, note: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=5)
    run = p.add_run(title)
    set_run_font(run, size=21, bold=True, color="000000")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=12)
    run = p.add_run(subtitle)
    set_run_font(run, size=10.5, color="000000")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14)
    run = p.add_run(note)
    set_run_font(run, size=9.5, color="000000")


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0)
    parse_inline(p, text, size=9.8, color="000000")


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [[item.strip() for item in line.strip().strip("|").split("|")] for line in lines if not re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", line)]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False

    if col_count == 2:
        width_ratios = [0.50, 0.50]
    elif col_count == 3:
        width_ratios = [0.33, 0.33, 0.34]
    elif col_count == 4 and rows[0][0] in {"顺序", "序号", "步骤"}:
        width_ratios = [0.08, 0.43, 0.21, 0.28]
    else:
        width_ratios = [1 / col_count] * col_count

    total_width_cm = 16.6
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is not None:
        tbl_width.set(qn("w:w"), str(int(Cm(total_width_cm).twips)))
        tbl_width.set(qn("w:type"), "dxa")

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx in range(col_count):
            value = values[idx] if idx < len(values) else ""
            cells[idx].width = Cm(total_width_cm * width_ratios[idx])
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.18)
            font_size = 8.1 if col_count == 4 else 8.8
            parse_inline(p, value, size=font_size, color="000000")
            if row_idx == 0:
                for run in p.runs:
                    run.bold = True
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def markdown_to_docx(source: Path, output: Path, title: str, subtitle: str, note: str, lanes: list[tuple[str, list[str]]]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in (("Title", 25, "000000"), ("Heading 1", 17, "000000"), ("Heading 2", 13, "000000"), ("Heading 3", 11.5, "000000")):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    add_cover(doc, title, subtitle, note)
    diagram_path = ASSET_DIR / f"{output.stem}_阶段泳道图.png"
    create_swimlane_diagram(diagram_path, title, lanes)
    body_started = False
    source_title_skipped = False
    diagram_inserted = False
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=2, after=8, line=1.1)
                p.paragraph_format.left_indent = Cm(0.55)
                p.paragraph_format.right_indent = Cm(0.35)
                for code_line in code_lines:
                    run = p.add_run(code_line + "\n")
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor.from_string("000000")
                in_code = False
                code_lines = []
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level, content = len(heading.group(1)), heading.group(2)
            if level == 1:
                if not source_title_skipped:
                    source_title_skipped = True
                    body_started = True
                    index += 1
                    continue
                if body_started:
                    doc.add_page_break()
                p = doc.add_paragraph(style="Heading 1")
                set_paragraph_spacing(p, before=8, after=10)
                p.paragraph_format.keep_with_next = True
                parse_inline(p, content, size=17, color="000000")
            else:
                style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[level]
                p = doc.add_paragraph(style=style)
                set_paragraph_spacing(p, before=13 if level == 2 else 8, after=5)
                p.paragraph_format.keep_with_next = True
                parse_inline(p, content, size={2: 15, 3: 12.5, 4: 11.2}[level], color="000000")
                if level == 2 and not diagram_inserted:
                    doc.add_picture(str(diagram_path), width=Inches(6.35))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.paragraphs[-1].paragraph_format.keep_with_next = True
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_spacing(caption, before=0, after=9)
                    run = caption.add_run("阶段泳道图")
                    set_run_font(run, size=8.5, color="000000")
                    diagram_inserted = True
            index += 1
            continue

        if not stripped:
            index += 1
            continue
        if stripped == "---":
            doc.add_paragraph()
            index += 1
            continue
        if stripped == ">":
            index += 1
            continue
        if stripped.startswith(">"):
            add_callout(doc, stripped.lstrip("> "))
            doc.add_paragraph()
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=3)
            p.paragraph_format.left_indent = Cm(0.62)
            p.paragraph_format.first_line_indent = Cm(-0.50)
            number_run = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(number_run, size=10.5)
            parse_inline(p, numbered.group(2), size=10.5)
        elif bullet:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=3)
            parse_inline(p, bullet.group(1), size=10.5)
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=6)
            parse_inline(p, stripped, size=10.5)
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for item in DOCS:
        item["output"].unlink(missing_ok=True)
        (ASSET_DIR / f"{item['output'].stem}_阶段泳道图.png").unlink(missing_ok=True)
        markdown_to_docx(item["source"], item["output"], item["title"], item["subtitle"], item["note"], item["lanes"])
    print("Created:")
    for item in DOCS:
        print(item["output"])


if __name__ == "__main__":
    main()
