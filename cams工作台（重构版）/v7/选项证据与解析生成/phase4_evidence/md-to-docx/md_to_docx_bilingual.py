# -*- coding: utf-8 -*-
"""
将小节 md 转为中英对照版 DOCX。

版式：
    1. 中文题干
       English stem
    A.中文选项
      English option
    答案:A
    解析:
    ...

说明：本脚本用于中英对照学习/教研复核版，不替代已验证的题库导入版。
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm
except ImportError:
    print("错误：缺少 python-docx，请先安装：pip install python-docx")
    sys.exit(1)

from md_to_docx import _normalize_import_option, _save_docx_with_retry
from md_to_docx_en import (
    _find_en_lookup,
    _load_en_lookup,
    add_label_paragraph,
    add_section_header,
    add_styled_paragraph,
    render_body_text,
    section_to_title,
)


def add_styled_paragraph_with_breaks(doc, lines, indent_cm=0, space_after=2):
    """添加一个 Word 段落，段内用软换行分隔多行，避免导入器把英文行识别成新选项。"""
    paragraph = add_styled_paragraph(doc, [], indent_cm=indent_cm, space_after=space_after)
    first_run = True
    for text in lines:
        if not first_run:
            paragraph.add_run().add_break()
        paragraph.add_run(str(text or ""))
        first_run = False
    return paragraph


def _split_question_blocks(content: str) -> list[str]:
    blocks = content.split("\n\n---\n\n")
    first_block = blocks[0].strip()
    header_end = first_block.find("\n教材章节：")
    if header_end > 0:
        first_q = first_block[header_end:].strip()
    else:
        first_q = ""
    question_blocks = [first_q] if first_q else []
    question_blocks += [b.strip() for b in blocks[1:] if b.strip()]
    return question_blocks


def _parse_question(block: str, en_lookup: dict) -> dict | None:
    """解析题块，保留中文题干/选项，并补充英文题干/选项。"""
    lines = block.strip().split("\n")
    result = {
        "stem": "",
        "stem_en": "",
        "options": [],
        "options_en": [],
        "answer": "",
        "analysis_text": "",
    }
    in_options = False
    opt_buf = None
    in_analysis = False
    sec_lines: list[str] = []
    qid = ""

    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith("教材章节："):
            continue
        if s.startswith("## "):
            qid = s[3:].strip()
            continue
        if s.startswith("题型："):
            continue
        if s.startswith("英文题干："):
            result["stem_en"] = s[5:].strip()
            continue
        if s.startswith("题干：") or s.startswith("题目："):
            result["stem"] = s[3:].strip()
            continue
        if s.startswith("English:"):
            if in_options and opt_buf:
                result["options_en"].append(s[len("English:"):].strip())
            elif not in_options:
                result["stem_en"] = s[len("English:"):].strip()
            continue
        if s == "选项：":
            in_options = True
            continue
        if in_options and s.startswith("- "):
            if opt_buf:
                result["options"].append(opt_buf)
            opt_buf = s[2:].strip()
            continue
        if s.startswith("答案："):
            in_options = False
            if opt_buf:
                result["options"].append(opt_buf)
                opt_buf = None
            result["answer"] = s[3:].strip()
            continue
        if s == "解析：":
            in_analysis = True
            in_options = False
            continue
        if in_analysis and s:
            if s == "---":
                break
            sec_lines.append(s)

    while len(result["options_en"]) < len(result["options"]):
        result["options_en"].append("")

    lookup = _find_en_lookup(result, qid, en_lookup)
    if lookup:
        if lookup.get("stem_en"):
            result["stem_en"] = lookup["stem_en"]
        if lookup.get("options_en"):
            json_opts = lookup["options_en"]
            result["options_en"] = [
                json_opts.get(
                    str(chr(65 + i)),
                    result["options_en"][i] if i < len(result["options_en"]) else "",
                )
                for i in range(len(result["options"]))
            ]

    result["analysis_text"] = "\n".join(sec_lines)
    return result if result["stem"] else None


def _strip_english_option_label(text: str, label: str) -> str:
    """去掉英文选项里可能自带的 A. / A、前缀，避免同一选项段内重复字母。"""
    text = str(text or "").strip()
    if not text:
        return ""
    return re.sub(rf"^{re.escape(label)}\s*[\.\、．]\s*", "", text, flags=re.IGNORECASE)


def _render_analysis(doc, analysis_text: str) -> None:
    add_label_paragraph(doc, "解析:", space_after=2)
    label_map = {
        "考点：": "考点",
        "核心解析：": "核心解析",
        "易错提醒：": "易错提醒",
        "教材原句：": "教材原句",
    }
    error_label_re = re.compile(r"^([A-H])项(\S+)：")
    for line in analysis_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        error_match = error_label_re.match(line)
        if error_match:
            label = f"{error_match.group(1)}项{error_match.group(2)}："
            add_section_header(doc, label)
            rest = line[len(label):].strip()
            if rest:
                render_body_text(doc, rest)
            continue
        matched = False
        for source_label, title in label_map.items():
            if line.startswith(source_label):
                add_section_header(doc, f"{title}：")
                rest = line[len(source_label):].strip()
                if rest:
                    render_body_text(doc, rest)
                matched = True
                break
        if not matched:
            render_body_text(doc, line)


def convert_md_to_docx(input_path: Path, output_path: Path) -> tuple[int, int]:
    content = input_path.read_text(encoding="utf-8")
    content = content.replace("「", "“").replace("」", "”")
    title = section_to_title(input_path.name)
    en_lookup = _load_en_lookup()

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    add_styled_paragraph(doc, [(f"试卷名称:{title} (中英对照)", False, False)], space_after=12)
    add_styled_paragraph(doc, [("一、单项选择题", False, False)], space_after=8)

    question_blocks = _split_question_blocks(content)
    total = 0
    for qi, block in enumerate(question_blocks, 1):
        q = _parse_question(block, en_lookup)
        if not q:
            continue
        total += 1

        stem_lines = [f"{qi}. {q['stem']}"]
        if q.get("stem_en"):
            stem_lines.append(q["stem_en"])
        add_styled_paragraph_with_breaks(doc, stem_lines, space_after=2)

        for index, option in enumerate(q["options"]):
            label = chr(65 + index)
            zh_option = _normalize_import_option(option)
            en_option = q["options_en"][index] if index < len(q["options_en"]) else ""
            option_lines = [zh_option]
            if en_option:
                en_option = _strip_english_option_label(en_option, label)
                option_lines.append(en_option)
            add_styled_paragraph_with_breaks(doc, option_lines, space_after=1)

        add_styled_paragraph(doc, [(f"答案:{q['answer']}", False, False)], space_after=2)
        _render_analysis(doc, q.get("analysis_text", ""))

        if qi < len(question_blocks):
            add_styled_paragraph(doc, [("", False, False)], space_after=6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_with_retry(doc, output_path)
    return total, 0


def main() -> None:
    parser = argparse.ArgumentParser(description="将小节 md 转为中英对照版 DOCX")
    parser.add_argument("-i", "--input", default="")
    parser.add_argument("-o", "--output", default="")
    parser.add_argument("--batch", default="", help="批量模式：指定目录，转换所有 p*-ch*-h*.md")
    args = parser.parse_args()

    if args.batch:
        batch_dir = Path(args.batch)
        if not batch_dir.exists():
            print(f"错误：目录不存在：{batch_dir}")
            sys.exit(1)
        md_files = sorted(batch_dir.glob("p*-ch*-h*.md"))
        if not md_files:
            print(f"错误：目录下无 p*-ch*-h*.md 文件：{batch_dir}")
            sys.exit(1)
        docx_dir = Path(__file__).resolve().parents[1] / "output" / "docx_bilingual"
        docx_dir.mkdir(parents=True, exist_ok=True)
        grand_total = 0
        for index, md_path in enumerate(md_files, 1):
            docx_path = docx_dir / md_path.with_suffix(".docx").name
            total, _ = convert_md_to_docx(md_path, docx_path)
            grand_total += total
            print(f"  [{index}/{len(md_files)}] {md_path.name} → {docx_path.name} ({total}题)")
        print(f"批量完成：{len(md_files)}文件, {grand_total}题, 输出到 {docx_dir}")
        return

    if not args.input or not args.output:
        print("错误：单文件模式需要 -i 和 -o，或使用 --batch 批量模式")
        sys.exit(1)

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        print(f"错误：输入文件不存在：{input_path}")
        sys.exit(1)
    total, _ = convert_md_to_docx(input_path, output_path)
    print(f"已保存：{output_path}（共{total}题）")


if __name__ == "__main__":
    main()
