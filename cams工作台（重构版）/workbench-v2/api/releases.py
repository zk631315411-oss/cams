from __future__ import annotations

import hashlib
import json
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.orm import Session

from .config import settings
from .content import version_markdown
from .db import Question, Release, ReleaseItem, Version


def create_release(db: Session, actor_id: int, title: str, question_ids: list[str]) -> Release:
    unique_ids = list(dict.fromkeys(question_ids))
    questions = list(db.scalars(select(Question).where(Question.id.in_(unique_ids))))
    found = {item.id: item for item in questions}
    missing = [item for item in unique_ids if item not in found]
    if missing:
        raise ValueError("题目不存在：" + ", ".join(missing))
    invalid = [item.id for item in questions if item.status != "approved"]
    if invalid:
        raise ValueError("仅可发布已批准题目：" + ", ".join(invalid))
    release_id = f"rel_{uuid.uuid4().hex}"
    release_dir = settings.content_root / "releases" / release_id
    release_dir.mkdir(parents=True, exist_ok=False)
    release = Release(id=release_id, created_by=actor_id, title=title, manifest_path=str(release_dir / "manifest.json"))
    db.add(release)
    manifest_items = []
    for question_id in unique_ids:
        question = found[question_id]
        version = db.get(Version, question.current_version_id)
        item = ReleaseItem(
            release_id=release_id,
            question_id=question.id,
            version_id=version.id,
            content_hash=version.content_hash,
        )
        db.add(item)
        manifest_items.append(
            {"question_id": question.id, "version_id": version.id, "content_hash": version.content_hash}
        )
    manifest = {"schema_version": "cams-workbench-release/v1", "release_id": release_id, "title": title, "items": manifest_items}
    Path(release.manifest_path).write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    db.commit()
    db.refresh(release)
    return release


def export_release_docx(db: Session, release: Release) -> Path:
    items = list(
        db.scalars(select(ReleaseItem).where(ReleaseItem.release_id == release.id).order_by(ReleaseItem.id))
    )
    document = Document()
    document.add_heading(release.title, level=1)
    for ordinal, item in enumerate(items, start=1):
        version = db.get(Version, item.version_id)
        parsed = version.parsed_content
        document.add_paragraph(f"{ordinal}. {parsed.get('stem_zh', '')}")
        for option in parsed.get("options", []):
            document.add_paragraph(f"{option['label']}.{option.get('zh', '')}")
        document.add_paragraph("答案:" + "、".join(parsed.get("answer_letters", [])))
        document.add_paragraph("解析:")
        for label, field in (
            ("考点", "exam_point"),
            ("核心解析", "core_analysis"),
            ("错误项分析", "wrong_analysis"),
            ("易错提醒", "reminder"),
        ):
            value = parsed.get(field, "")
            if value:
                heading = document.add_paragraph()
                run = heading.add_run(label + "：")
                run.bold = True
                run.font.size = Pt(10.5)
                document.add_paragraph(value)
        document.add_paragraph("")
    release_dir = Path(release.manifest_path).parent
    output = release_dir / f"{release.id}.docx"
    document.save(output)
    release.export_path = str(output)
    release.export_hash = hashlib.sha256(output.read_bytes()).hexdigest()
    release.state = "exported"
    db.commit()
    return output
