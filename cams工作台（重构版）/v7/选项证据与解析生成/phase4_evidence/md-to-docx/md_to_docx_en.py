# -*- coding: utf-8 -*-
"""
将润色后的小节 md 转为纯英文版 DOCX（题干+选项英文，解析中文）。

单文件：
    python md_to_docx_en.py -i sections/p1-ch1-h2.md -o docx_en/p1-ch1-h2.docx

批量：
    python md_to_docx_en.py --batch sections/
"""

import argparse
import difflib
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    print("错误：缺少 python-docx，请先安装：pip install python-docx")
    sys.exit(1)

# <sup>PXXX</sup> 和 **bold** 正则
_SUP_RE = re.compile(r"<sup>(P\d+)</sup>|（(书内第\d+页|[Pp]\d+[、，,Pp\d]*)）")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# 试卷名称映射
_TITLE_RE = re.compile(r"p(\d+)-ch(\d+)-h(\d+)")

# v7_questions.json 路径
_QUESTIONS_JSON_PATH = Path(__file__).resolve().parents[2] / "phase3.5_questions" / "output" / "v7_questions.json"
_EXPLANATIONS_DIR = Path(__file__).resolve().parents[1] / "output" / "explanations"

# 英文数据缓存
_EN_LOOKUP: dict | None = None


def _normalize_for_match(text: str) -> str:
    """用于中英文源题回填时的宽松匹配：去掉空白、标点和题号差异。"""
    text = str(text or "").lower()
    text = re.sub(r"^\[[^\]]+\]", "", text)
    text = re.sub(r"^[a-e]\s*[\.\、．]\s*", "", text)
    return re.sub(r"[\s，,。\.？?！!：:；;、（）()\[\]【】“”\"'‘’\-—_]+", "", text)


def _strip_option_label(text: str) -> str:
    return re.sub(r"^[A-Ea-e]\s*[\.\、．]\s*", "", str(text or "").strip())


def _load_en_lookup() -> dict:
    """加载 v7_questions.json，建立题号、中文题干和题干+选项的英文回填索引。"""
    global _EN_LOOKUP
    if _EN_LOOKUP is not None:
        return _EN_LOOKUP
    _EN_LOOKUP = {"by_qid": {}, "by_stem": {}, "items": []}
    if not _QUESTIONS_JSON_PATH.exists():
        print(f"  警告：{_QUESTIONS_JSON_PATH} 不存在，英文回退到 MD 内联提取")
        return _EN_LOOKUP
    with open(_QUESTIONS_JSON_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    for item in data.get("items", []) or []:
        qid = item.get("question_id", "")
        stem_zh = str(item.get("stem", "") or "").strip()
        options_zh = item.get("options", {}) or {}
        options_en = item.get("options_en", {}) or {}
        lookup = {
            "stem_en": str(item.get("stem_en", "") or "").strip(),
            "options_en": {str(k): str(v or "") for k, v in options_en.items()},
            "stem_key": _normalize_for_match(stem_zh),
            "options_key": _normalize_for_match(" ".join(str(options_zh.get(k, "") or "") for k in sorted(options_zh))),
        }
        if qid:
            _EN_LOOKUP["by_qid"][qid] = lookup
        if lookup["stem_key"]:
            _EN_LOOKUP["by_stem"][lookup["stem_key"]] = lookup
        _EN_LOOKUP["items"].append(lookup)

    # 当前小节 MD 由 output/explanations/v7_q_*.md 生成；这些源文件的文件名带题号，
    # 且内容是润色后的中文题干/选项。用它们补一层索引，避免润色后中文与原始 JSON 差异过大时回填失败。
    if _EXPLANATIONS_DIR.exists():
        for exp_path in sorted(_EXPLANATIONS_DIR.glob("v7_q_*.md")):
            qid = exp_path.stem
            lookup = _EN_LOOKUP["by_qid"].get(qid)
            if not lookup:
                continue
            try:
                exp_text = exp_path.read_text(encoding="utf-8")
            except OSError:
                continue
            stem_m = re.search(r"^题干：(.+)$", exp_text, flags=re.MULTILINE)
            stem_en_m = re.search(r"^英文题干：(.+)$", exp_text, flags=re.MULTILINE)
            if stem_en_m:
                # 解析 md 可能经过人工修订；英文导出应优先使用最新版 md 中的英文题干，
                # 避免被原始题库 JSON 的旧英文覆盖。
                lookup["stem_en"] = stem_en_m.group(1).strip()
            for option_en_m in re.finditer(
                r"^-\s+([A-Ha-h])\s*[\.\、．].*?\n\s*English:\s*(.+)$",
                exp_text,
                flags=re.MULTILINE,
            ):
                label = option_en_m.group(1).upper()
                # 同上，选项英文以最新版解析 md 为准。
                lookup["options_en"][label] = option_en_m.group(2).strip()
            if stem_m:
                exp_stem_key = _normalize_for_match(stem_m.group(1))
                if exp_stem_key:
                    _EN_LOOKUP["by_stem"][exp_stem_key] = lookup
                    _EN_LOOKUP["items"].append({
                        **lookup,
                        "stem_key": exp_stem_key,
                        "options_key": lookup.get("options_key", ""),
                    })
            option_lines = re.findall(r"^-\s+[A-Ea-e]\s*[\.\、．]\s*(.+)$", exp_text, flags=re.MULTILINE)
            if option_lines:
                exp_options_key = _normalize_for_match(" ".join(option_lines))
                if exp_options_key:
                    _EN_LOOKUP["items"].append({
                        **lookup,
                        "stem_key": exp_stem_key if stem_m else lookup.get("stem_key", ""),
                        "options_key": exp_options_key,
                    })
    return _EN_LOOKUP


def _find_en_lookup(result: dict, qid: str, en_lookup: dict) -> dict | None:
    """优先按题号，其次按中文题干精确归一化，最后按题干+选项相似度兜底。"""
    if qid and qid in en_lookup.get("by_qid", {}):
        return en_lookup["by_qid"][qid]

    stem_key = _normalize_for_match(result.get("stem", ""))
    if stem_key and stem_key in en_lookup.get("by_stem", {}):
        return en_lookup["by_stem"][stem_key]

    option_key = _normalize_for_match(" ".join(_strip_option_label(o) for o in result.get("options", [])))
    if not stem_key and not option_key:
        return None

    best: tuple[float, float, float, dict] | None = None
    for item in en_lookup.get("items", []):
        stem_ratio = difflib.SequenceMatcher(None, stem_key, item.get("stem_key", "")).ratio() if stem_key else 0.0
        option_ratio = difflib.SequenceMatcher(None, option_key, item.get("options_key", "")).ratio() if option_key else 0.0
        score = max(stem_ratio, stem_ratio * 0.35 + option_ratio * 0.65, option_ratio * 0.92)
        if best is None or score > best[0]:
            best = (score, stem_ratio, option_ratio, item)

    if not best:
        return None
    score, stem_ratio, option_ratio, item = best
    if score >= 0.84 or (score >= 0.78 and stem_ratio >= 0.50 and option_ratio >= 0.65):
        return item
    return None


def section_to_title(filename: str) -> str:
    m = _TITLE_RE.search(filename)
    if not m:
        return "CAMS"
    return f"CAMS CH{m.group(2).zfill(2)}"


def _set_run_bg(run, color: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:fill"), color)
    shd.set(_qn("w:val"), "clear")
    run._element.get_or_add_rPr().append(shd)


def add_styled_paragraph(doc, segments, indent_cm=0, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    for text, bold, is_sup in segments:
        run = p.add_run(text)
        if bold:
            run.font.bold = True
        if is_sup:
            run.font.superscript = True
            run.font.size = Pt(7.5)
    return p


def add_section_header(doc, label: str):
    """添加蓝色文字底色的解析小节标题，如「考点」「核心解析」。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(10.5)
    _set_run_bg(run, "D6E4F0")
    return p


def add_label_paragraph(doc, label: str, space_after=2):
    """添加结构标签段落，仅给标签文字加蓝底和加粗。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(label)
    run.font.bold = True
    _set_run_bg(run, "D6E4F0")
    return p


def render_section_label(doc, label: str):
    """渲染加粗的段落标签，如「考点：」「核心解析：」等。"""
    add_label_paragraph(doc, label, space_after=0)


def parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    tokens = []
    pos = 0
    while pos < len(text):
        sup_m = _SUP_RE.search(text, pos)
        bold_m = _BOLD_RE.search(text, pos)
        next_m = None
        tag_type = None
        if sup_m and (not bold_m or sup_m.start() <= bold_m.start()):
            next_m = sup_m
            tag_type = "sup"
        elif bold_m:
            next_m = bold_m
            tag_type = "bold"
        if not next_m:
            tokens.append((text[pos:], False, False))
            break
        if next_m.start() > pos:
            tokens.append((text[pos:next_m.start()], False, False))
        if tag_type == "sup":
            label = (next_m.group(1) or next_m.group(2) or "").strip()
            tokens.append((label, False, True))
        else:
            tokens.append((next_m.group(1), True, False))
        pos = next_m.end()
    return tokens


def render_body_text(doc, text: str):
    text = text.strip()
    if not text:
        return
    segments = parse_inline(text)
    add_styled_paragraph(doc, segments, indent_cm=0.74)


def convert_md_to_docx(input_path: Path, output_path: Path) -> tuple[int, int]:
    content = input_path.read_text(encoding="utf-8")
    content = content.replace("「", "“").replace("」", "”")
    title = section_to_title(input_path.name)
    en_lookup = _load_en_lookup()

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)

    add_styled_paragraph(doc, [(f"试卷名称:{title} (EN)", False, False)], space_after=12)
    add_styled_paragraph(doc, [("一、单项选择题", False, False)], space_after=8)

    blocks = content.split("\n\n---\n\n")
    first_block = blocks[0].strip()
    header_end = first_block.find("\n教材章节：")
    if header_end > 0:
        first_q = first_block[header_end:].strip()
    else:
        first_q = ""
    question_blocks = [first_q] if first_q else []
    question_blocks += [b.strip() for b in blocks[1:] if b.strip()]

    total = 0
    for qi, block in enumerate(question_blocks, 1):
        q = _parse_question(block, en_lookup)
        if not q:
            continue
        total += 1

        # 英文题目
        add_styled_paragraph(doc, [(f"{qi}. {q['stem']}", False, False)], space_after=2)

        # 英文选项（带字母前缀）
        for idx, opt in enumerate(q["options"]):
            label = chr(65 + idx)  # A, B, C, D...
            add_styled_paragraph(doc, [(f"{label}. {opt}", False, False)], indent_cm=0.5, space_after=1)

        # 答案
        add_styled_paragraph(doc, [(f"答案：{q['answer']}", False, False)], space_after=2)

        # 解析（中文，蓝底标题 + 内容）
        add_label_paragraph(doc, "解析：", space_after=2)
        _LABEL_MAP = {
            "考点：": "考点", "核心解析：": "核心解析",
            "易错提醒：": "易错提醒", "教材原句：": "教材原句",
        }
        _ERROR_LABEL_RE = re.compile(r"^([A-E])项(\S+)：")
        for line in q.get("analysis_text", "").split("\n"):
            line = line.strip()
            if not line:
                continue
            em = _ERROR_LABEL_RE.match(line)
            if em:
                add_section_header(doc, f"{em.group(1)}项{em.group(2)}：")
                rest = line[len(f"{em.group(1)}项{em.group(2)}："):]
                render_body_text(doc, rest)
                continue
            matched = False
            for lbl, title in _LABEL_MAP.items():
                if line.startswith(lbl):
                    add_section_header(doc, f"{title}：")
                    rest = line[len(lbl):]
                    if rest:
                        render_body_text(doc, rest)
                    matched = True
                    break
            if not matched:
                render_body_text(doc, line)

        if qi < len(question_blocks):
            add_styled_paragraph(doc, [("", False, False)], space_after=6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return total, 0


def _parse_question(block: str, en_lookup: dict) -> dict | None:
    """解析题块，优先用 JSON 查英文，回退 MD 内联提取。"""
    lines = block.strip().split("\n")
    result = {
        "stem": "", "stem_en": "", "options": [], "options_en": [],
        "answer": "", "analysis_text": "",
    }
    in_options = False
    opt_buf = None
    in_analysis = False
    sec_lines: list[str] = []
    qid: str = ""

    for line in lines:
        s = line.strip()

        if s.startswith("教材章节："):
            continue
        if s.startswith("## "):
            qid = s[3:].strip()
            continue
        if s.startswith("题型："):
            continue
        elif s.startswith("英文题干："):
            result["stem_en"] = s[5:].strip()
            continue
        elif s.startswith("题干：") or s.startswith("题目："):
            result["stem"] = s[3:].strip()
        elif s.startswith("English:"):
            if in_options and opt_buf:
                # 选项级 English
                result["options_en"].append(s[len("English:"):].strip())
            elif not in_options:
                # 新格式：题目后的 stem_en
                result["stem_en"] = s[len("English:"):].strip()
            continue
        elif s == "选项：":
            in_options = True
        elif in_options and s.startswith("- "):
            if opt_buf:
                result["options"].append(opt_buf)
            opt_buf = s[2:].strip()
        elif s.startswith("答案："):
            in_options = False
            if opt_buf:
                result["options"].append(opt_buf)
                opt_buf = None
            result["answer"] = s[3:].strip()
        elif s == "解析：":
            in_analysis = True
            in_options = False
        elif in_analysis and s:
            if s == "---":
                break
            sec_lines.append(s)

    # 对补 option_en 长度（与 options 对齐）
    while len(result["options_en"]) < len(result["options"]):
        result["options_en"].append("")

    # JSON 查英文（优先题号；小节 MD 没有题号时按中文题干/选项兜底匹配）
    lookup = _find_en_lookup(result, qid, en_lookup)
    if lookup:
        if lookup.get("stem_en"):
            result["stem_en"] = lookup["stem_en"]
        if lookup.get("options_en"):
            json_opts = lookup["options_en"]
            result["options_en"] = [
                json_opts.get(str(chr(65 + i)), result["options_en"][i] if i < len(result["options_en"]) else "")
                for i in range(len(result["options"]))
            ]

    # 用英文 stem / options 替换
    if result["stem_en"]:
        result["stem"] = result["stem_en"]
    if result["options_en"] and any(result["options_en"]):
        result["options"] = result["options_en"]

    result["analysis_text"] = "\n".join(sec_lines)
    return result if result["stem"] else None


def main():
    parser = argparse.ArgumentParser(description="将小节 md 转为纯英文版 DOCX")
    parser.add_argument("-i", "--input", default="")
    parser.add_argument("-o", "--output", default="")
    parser.add_argument("--batch", default="", help="批量模式：指定目录，转换所有 p*-ch*-h*.md")
    args = parser.parse_args()

    if args.batch:
        batch_dir = Path(args.batch)
        if not batch_dir.is_dir():
            print(f"错误：目录不存在：{batch_dir}")
            sys.exit(1)
        md_files = sorted(batch_dir.glob("p*-ch*-h*.md"))
        if not md_files:
            print(f"错误：目录下无 p*-ch*-h*.md 文件：{batch_dir}")
            sys.exit(1)
        # 统一输出到 phase4_evidence/output/docx_en
        docx_dir = Path(__file__).resolve().parents[1] / "output" / "docx_en"
        docx_dir.mkdir(parents=True, exist_ok=True)
        grand_total = 0
        for i, md_path in enumerate(md_files, 1):
            docx_path = docx_dir / md_path.with_suffix(".docx").name
            total, _ = convert_md_to_docx(md_path, docx_path)
            grand_total += total
            print(f"  [{i}/{len(md_files)}] {md_path.name} → {docx_path.name} ({total}题)")
        print(f"批量完成：{len(md_files)}文件, {grand_total}题, 输出到 {docx_dir}")
        return

    if not args.input or not args.output:
        print("错误：单文件模式需要 -i 和 -o，或使用 --batch 批量模式")
        sys.exit(1)

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        print(f"错误：输入文件不存在：{input_path}")
        sys.exit(1)

    total, _ = convert_md_to_docx(input_path, output_path)
    print(f"已保存：{output_path}（共{total}题）")


if __name__ == "__main__":
    main()
