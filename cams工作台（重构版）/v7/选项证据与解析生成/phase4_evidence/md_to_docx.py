# -*- coding: utf-8 -*-
"""
将 software_export/chapters/CHxx.md 转为结构化 docx。
支持朴素字体排版：黑体标题 + 宋体正文 + Times New Roman 英文。

用法示例：
    python md_to_docx.py \
        -i "D:\...\software_export\chapters\CH02.md" \
        -o "D:\...\software_export\chapters\CH02.docx"
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError as e:
    print("错误：缺少 python-docx，请先安装：pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name='宋体', font_size=10.5, bold=False, italic=False):
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(doc, text, font_name='宋体', font_size=10.5, bold=False, italic=False,
                  align=None, first_indent=None, left_indent=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name, font_size, bold, italic)
    if align:
        p.alignment = align
    if first_indent is not None:
        p.paragraph_format.first_line_indent = first_indent
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p


def parse_question_block(block):
    lines = block.strip().split('\n')
    result = {
        'id': '',
        'type': '',
        'question_cn': '',
        'question_en': '',
        'options': [],
        'answer': '',
        'sections': {}
    }
    current_section = None
    section_buffer = []
    in_options = False
    option_buffer = None
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        if stripped.startswith('## '):
            result['id'] = stripped.replace('## ', '').strip()
        elif stripped.startswith('题型：'):
            result['type'] = stripped.replace('题型：', '').strip()
        elif stripped.startswith('题目：'):
            result['question_cn'] = stripped.replace('题目：', '').strip()
        elif stripped.startswith('English:') and not in_options:
            result['question_en'] = stripped.replace('English:', '').strip()
        elif stripped == '选项：':
            in_options = True
        elif in_options and stripped.startswith('- '):
            if option_buffer:
                result['options'].append(option_buffer)
            option_text = stripped[2:].strip()
            option_buffer = {'text': option_text, 'english': ''}
        elif in_options and stripped.startswith('English:') and option_buffer:
            option_buffer['english'] = stripped.replace('English:', '').strip()
        elif stripped.startswith('答案：'):
            in_options = False
            if option_buffer:
                result['options'].append(option_buffer)
                option_buffer = None
            result['answer'] = stripped.replace('答案：', '').strip()
        elif stripped == '解析：':
            in_options = False
            if option_buffer:
                result['options'].append(option_buffer)
                option_buffer = None
            current_section = None
        elif stripped.startswith('【') and stripped.endswith('】'):
            if current_section and section_buffer:
                result['sections'][current_section] = '\n'.join(section_buffer).strip()
                section_buffer = []
            current_section = stripped[1:-1]
        elif current_section and line:
            section_buffer.append(line)

        i += 1

    if current_section and section_buffer:
        result['sections'][current_section] = '\n'.join(section_buffer).strip()
    return result


def add_question_to_doc(doc, q):
    add_paragraph(doc, q['id'], font_name='黑体', font_size=12, bold=True, space_before=12, space_after=6)

    p_type = doc.add_paragraph()
    run_type_label = p_type.add_run('题型：')
    set_run_font(run_type_label, '宋体', 10.5, bold=True)
    run_type_val = p_type.add_run(q['type'])
    set_run_font(run_type_val, '宋体', 10.5, bold=True)
    p_type.paragraph_format.space_after = Pt(4)

    p_q = doc.add_paragraph()
    run_q_label = p_q.add_run('题目：')
    set_run_font(run_q_label, '宋体', 10.5, bold=True)
    run_q_text = p_q.add_run(q['question_cn'])
    set_run_font(run_q_text, '宋体', 10.5, bold=True)
    p_q.paragraph_format.space_after = Pt(2)

    if q['question_en']:
        add_paragraph(doc, q['question_en'], font_name='Times New Roman', font_size=10, italic=True,
                      left_indent=Cm(0.5), space_after=4)

    add_paragraph(doc, '选项：', font_name='宋体', font_size=10.5, bold=True, space_after=2)

    correct_set = {c for c in q['answer'] if c in 'ABCDE'}

    for opt in q['options']:
        p_opt = doc.add_paragraph()
        p_opt.paragraph_format.left_indent = Cm(0.5)
        p_opt.paragraph_format.space_after = Pt(1)
        opt_letter = opt['text'][0] if opt['text'] else ''
        is_correct = opt_letter in correct_set
        run_letter = p_opt.add_run(opt['text'])
        if is_correct:
            set_run_font(run_letter, '宋体', 10.5, bold=True)
        else:
            set_run_font(run_letter, '宋体', 10.5)
        if opt['english']:
            p_opt_en = doc.add_paragraph()
            p_opt_en.paragraph_format.left_indent = Cm(0.8)
            p_opt_en.paragraph_format.space_after = Pt(2)
            run_en = p_opt_en.add_run(opt['english'])
            set_run_font(run_en, 'Times New Roman', 10, italic=True)

    add_paragraph(doc, f'答案：{q["answer"]}', font_name='宋体', font_size=10.5, bold=True, space_after=6)

    section_order = ['考点', '核心解析', '选项分析', '易错提醒']
    for sec_name in section_order:
        if sec_name not in q['sections']:
            continue
        add_paragraph(doc, f'【{sec_name}】', font_name='黑体', font_size=10.5, bold=True,
                      space_before=8, space_after=3)
        content = q['sections'][sec_name]
        if sec_name == '选项分析':
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                m = re.match(r'^([A-E]项(?:正确|错误)[（(][^)）]*[）)])', line)
                if m:
                    prefix = m.group(1)
                    rest = line[len(prefix):]
                    if rest.startswith('：'):
                        rest = rest[1:]
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    p.paragraph_format.space_after = Pt(2)
                    run_p = p.add_run(prefix)
                    set_run_font(run_p, '宋体', 10.5, bold=True)
                    run_r = p.add_run(rest)
                    set_run_font(run_r, '宋体', 10.5)
                else:
                    add_paragraph(doc, line, first_indent=Cm(0.74), space_after=2)
        else:
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    add_paragraph(doc, line, first_indent=Cm(0.74), space_after=2)

    doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser(description='将题库软件版解析预览 Markdown 转为结构化 docx')
    parser.add_argument('-i', '--input', required=True, help='输入 md 文件路径')
    parser.add_argument('-o', '--output', required=True, help='输出 docx 文件路径')
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        print(f'错误：输入文件不存在：{input_path}')
        sys.exit(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    parts = content.split('## ')
    header = parts[0]
    question_blocks = ['## ' + p for p in parts[1:]]

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    title_match = re.search(r'^# (.+)$', header, re.MULTILINE)
    if title_match:
        add_paragraph(doc, title_match.group(1), font_name='黑体', font_size=16, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    count_match = re.search(r'可导出题目数：(\d+)', header)
    if count_match:
        add_paragraph(doc, f'可导出题目数：{count_match.group(1)}', font_name='宋体', font_size=10.5,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    for block in question_blocks:
        q = parse_question_block(block)
        if q['id']:
            add_question_to_doc(doc, q)

    doc.save(output_path)
    print(f'已保存：{output_path}')


if __name__ == '__main__':
    main()
