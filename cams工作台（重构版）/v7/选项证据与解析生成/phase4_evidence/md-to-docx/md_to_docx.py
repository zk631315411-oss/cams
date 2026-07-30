# -*- coding: utf-8 -*-
"""
将润色后的小节 md 转为标准试卷 DOCX 格式。

单文件：
    python md_to_docx.py -i sections/p1-ch1-h2.md -o sections/p1-ch1-h2.docx

批量：
    python md_to_docx.py --batch sections/
"""

import argparse
import errno
import gc
import re
import sys
import time
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    print("错误：缺少 python-docx，请先安装：pip install python-docx")
    sys.exit(1)

# <sup>PXXX</sup>、括号页码、裸 PXXX 和 **bold** 正则
# 纯页码括号会被整体消费，DOCX只输出上标页码：
# <sup>P28</sup> / （书内第28页） / （P28） / （P28、P42） / P28
# 裸页码要求字母数字边界，避免把 P2P 等业务术语误识别为页码。
_PAGE_REF_PATTERN = (
    r"[Pp]\d+"
    r"(?:(?:\s*[-–—‑]\s*[Pp]?\d+)|(?:\s*[/、,，]\s*[Pp]\d+))*"
)
_SUP_RE = re.compile(
    rf"<sup>(?P<html>{_PAGE_REF_PATTERN})</sup>"
    r"|（(?P<book>书内第\d+页)）"
    rf"|[（(](?P<paren>{_PAGE_REF_PATTERN})[）)]"
    rf"|(?<![A-Za-z0-9])(?P<bare>{_PAGE_REF_PATTERN})(?![A-Za-z0-9])"
)
_PAGE_REF_RE = re.compile(rf"^{_PAGE_REF_PATTERN}$")
_WORD_JOINER = "\u2060"
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# 试卷名称映射：p1-ch1-h2 → CAMS CH01
_TITLE_RE = re.compile(r"p(\d+)-ch(\d+)-h(\d+)")
_IMPORT_QUESTION_RE = re.compile(r"^(\d+)\.\s+\S.+$")
_IMPORT_OPTION_RE = re.compile(r"^[A-H]\.\s*\S.+")
_IMPORT_ANSWER_RE = re.compile(r"^答案[:：][A-H](?:[、,，][A-H])*$")
_ANALYSIS_NUMBER_RE = re.compile(r"^(?:\d+[.．、]|[（(]\d+[)）])")
DEFAULT_TEMPLATE_PATH = Path.home() / "Downloads" / "试题模板.docx"


def section_to_title(filename: str) -> str:
    """从文件名提取试卷标题，如 p1-ch1-h2.md → CAMS CH01。"""
    m = _TITLE_RE.search(filename)
    if not m:
        return "CAMS"
    return f"CAMS CH{m.group(2).zfill(2)}"


def _set_run_bg(run, color: str):
    """给 run 设置文字底色。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn
    shd = OxmlElement("w:shd")
    shd.set(_qn("w:fill"), color)
    shd.set(_qn("w:val"), "clear")
    run._element.get_or_add_rPr().append(shd)


def add_styled_paragraph(doc, segments, indent_cm=0, space_after=2):
    """添加段落，segments 为 [(text, bold, superscript), ...]。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    for text, bold, is_sup in segments:
        run = p.add_run(text)
        if bold:
            run.font.bold = True
        if is_sup:
            run.font.superscript = True
            run.font.size = Pt(7.5)
    return p


def add_section_header(doc, label: str):
    """添加蓝色文字底色的解析小节标题，如「考点」「核心解析」。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    run.font.bold = True
    run.font.size = Pt(10.5)
    _set_run_bg(run, "D6E4F0")
    return p


def add_label_paragraph(doc, label: str, space_after=2):
    """添加题库导入可识别的标签段落，仅给标签文字加蓝底和加粗。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(label)
    run.font.bold = True
    _set_run_bg(run, "D6E4F0")
    return p


def _save_docx_with_retry(doc, output_path: Path, attempts: int = 5) -> None:
    """重试Windows连续覆盖DOCX时偶发的EINVAL写入冲突。"""
    for attempt in range(1, attempts + 1):
        try:
            doc.save(output_path)
            return
        except OSError as exc:
            if exc.errno != errno.EINVAL or attempt == attempts:
                raise
            gc.collect()
            time.sleep(0.2 * attempt)


def _keep_page_numbers_together(text: str) -> str:
    """将单页或复合页码作为一个不可断行单元。"""
    return text


def parse_inline(text: str) -> list[tuple[str, bool, bool]]:
    """解析行内 <sup> 和 **bold**，返回 [(text, bold, sup), ...] 列表。"""
    # 合并两个正则，按出现顺序解析
    tokens = []
    pos = 0
    while pos < len(text):
        sup_m = _SUP_RE.search(text, pos)
        bold_m = _BOLD_RE.search(text, pos)

        next_m = None
        tag_type = None
        if sup_m and (not bold_m or sup_m.start() <= bold_m.start()):
            next_m = sup_m
            tag_type = "sup"
        elif bold_m:
            next_m = bold_m
            tag_type = "bold"

        if not next_m:
            tokens.append((text[pos:], False, False))
            break

        if next_m.start() > pos:
            tokens.append((text[pos:next_m.start()], False, False))

        if tag_type == "sup":
            sup_text = next(
                (next_m.group(name) for name in ("html", "book", "paren", "bare")
                 if next_m.group(name)),
                "",
            ).strip()
            tokens.append((sup_text, False, True))
        else:
            tokens.append((next_m.group(1), True, False))

        pos = next_m.end()

    return tokens


def render_section_label(doc, label: str):
    """渲染加粗的段落标签，如「考点：」「核心解析：」等。"""
    add_label_paragraph(doc, label, space_after=0)


def render_body_text(doc, text: str):
    """渲染正文段落，支持行内上标和加粗。"""
    text = text.strip()
    if not text:
        return
    segments = parse_inline(text)
    add_styled_paragraph(doc, segments, indent_cm=0.74)


def _document_from_template(template_path: Path):
    """加载题库模板并清空示例正文，保留模板的样式和页面设置。"""
    if not template_path.exists():
        raise FileNotFoundError(f"题库模板不存在：{template_path}")

    from docx.oxml.ns import qn

    doc = Document(template_path)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return doc


def _normalize_import_option(option: str) -> str:
    """按已验证的导入格式输出 A.选项，标签后不留空格。"""
    return re.sub(r"^([A-H])\.\s*", r"\1.", option.strip())


def _prepare_analysis_lines(analysis_text: str) -> list[str]:
    """把解析区的编号项并入上一段，避免后台将其识别为新题。"""
    lines: list[str] = []
    for raw_line in analysis_text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if _ANALYSIS_NUMBER_RE.match(line):
            if lines:
                separator = " " if lines[-1].endswith(("：", ":")) else "；"
                lines[-1] += separator + line
            else:
                lines.append("要点：" + line)
            continue
        lines.append(line)
    return lines


def _render_import_analysis(doc, analysis_text: str) -> None:
    """按题库导入协议渲染解析，同时保留行内上标和加粗。"""
    lines = _prepare_analysis_lines(analysis_text)
    add_label_paragraph(doc, "解析:", space_after=2)
    if not lines:
        return

    error_heading_added = False
    option_analysis_re = re.compile(
        r"^([A-H](?:[、,，][A-H])*)项([^：]+)：(.*)$"
    )
    section_labels = ("考点：", "核心解析：", "教材原句：", "易错提醒：")
    for line in lines:
        section_label = next(
            (label for label in section_labels if line.startswith(label)),
            None,
        )
        if section_label:
            add_label_paragraph(doc, section_label, space_after=2)
            body = line[len(section_label):].strip()
            if body:
                add_styled_paragraph(doc, parse_inline(body), space_after=2)
            continue

        option_match = option_analysis_re.match(line)
        if option_match:
            if not error_heading_added:
                add_label_paragraph(doc, "错误项分析：", space_after=2)
                error_heading_added = True
            label = f"{option_match.group(1)}项{option_match.group(2)}："
            add_label_paragraph(doc, label, space_after=2)
            body = option_match.group(3).strip()
            if body:
                add_styled_paragraph(doc, parse_inline(body), space_after=2)
            continue
        add_styled_paragraph(doc, parse_inline(line), space_after=2)


def _validate_import_docx(output_path: Path, expected_total: int) -> None:
    """按题库模板的文本协议回读DOCX，阻止结构不完整的文件发布。"""
    doc = Document(output_path)
    lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    question_indexes = []
    for index, line in enumerate(lines):
        match = _IMPORT_QUESTION_RE.match(line)
        if not match:
            continue
        lookahead = lines[index + 1 : index + 14]
        option_count = sum(1 for item in lookahead if _IMPORT_OPTION_RE.match(item))
        has_answer = any(_IMPORT_ANSWER_RE.match(item) for item in lookahead)
        if option_count >= 2 and has_answer:
            question_indexes.append(index)
    if len(question_indexes) != expected_total:
        raise ValueError(
            f"{output_path.name}: 题目识别数 {len(question_indexes)} != {expected_total}"
        )

    for question_number, start in enumerate(question_indexes, 1):
        end = (
            question_indexes[question_number]
            if question_number < len(question_indexes)
            else len(lines)
        )
        block = lines[start:end]
        match = _IMPORT_QUESTION_RE.match(block[0])
        if not match or int(match.group(1)) != question_number:
            raise ValueError(f"{output_path.name}: 第{question_number}题题号不连续")

        option_count = sum(bool(_IMPORT_OPTION_RE.match(line)) for line in block)
        answer_count = sum(bool(_IMPORT_ANSWER_RE.match(line)) for line in block)
        analysis_count = sum(line.startswith("解析:") for line in block)
        if option_count < 2:
            raise ValueError(f"{output_path.name}: 第{question_number}题选项不足")
        if answer_count != 1:
            raise ValueError(f"{output_path.name}: 第{question_number}题答案格式异常")
        if analysis_count != 1:
            raise ValueError(f"{output_path.name}: 第{question_number}题解析标记异常")

        for line in block[1:]:
            if _ANALYSIS_NUMBER_RE.match(line):
                raise ValueError(
                    f"{output_path.name}: 第{question_number}题存在未转义行首编号: {line}"
                )


def convert_md_to_docx(
    input_path: Path,
    output_path: Path,
    template_path: Path = DEFAULT_TEMPLATE_PATH,
) -> tuple[int, int]:
    content = input_path.read_text(encoding="utf-8")
    # 中文弯引号替换为直引号
    content = content.replace("「", "“").replace("」", "”")
    title = section_to_title(input_path.name)

    doc = _document_from_template(template_path)

    # 试卷名称
    add_styled_paragraph(doc, [(f"试卷名称:{title}", False, False)], space_after=12)

    # 一、单项选择题
    add_styled_paragraph(doc, [("一、单项选择题", False, False)], space_after=8)

    # 切分题目：以 --- 为分隔符
    blocks = content.split("\n\n---\n\n")
    # 第一块是 section header + 第一题，去掉 section header 得到第一题
    first_block = blocks[0].strip()
    header_end = first_block.find("\n教材章节：")
    if header_end > 0:
        first_q = first_block[header_end:].strip()
    else:
        first_q = ""
    # 后续各题
    question_blocks = [first_q] if first_q else []
    question_blocks += [b.strip() for b in blocks[1:] if b.strip()]

    total = 0
    for qi, block in enumerate(question_blocks, 1):
        q = _parse_question(block)
        if not q:
            continue
        total += 1

        # 题目
        add_styled_paragraph(doc, [(f"{qi}. {q['stem']}", False, False)], space_after=2)

        # 选项 A-D
        for opt in q["options"]:
            normalized_option = _normalize_import_option(opt)
            add_styled_paragraph(doc, [(normalized_option, False, False)], space_after=1)

        # 答案
        add_styled_paragraph(doc, [(f"答案:{q['answer']}", False, False)], space_after=2)

        _render_import_analysis(doc, q.get("analysis_text", ""))

        if qi < len(question_blocks):
            add_styled_paragraph(doc, [("", False, False)], space_after=6)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    _save_docx_with_retry(doc, output_path)
    del doc
    _validate_import_docx(output_path, total)
    gc.collect()
    return total, 0


def _parse_question(block: str) -> dict | None:
    """解析单个题块，提取题干、选项、答案、解析。"""
    lines = block.strip().split("\n")
    result = {"stem": "", "options": [], "answer": "", "analysis": {}}
    in_options = False
    opt_buf = None
    in_analysis = False
    current_sec = None
    sec_lines: list[str] = []

    for line in lines:
        s = line.strip()

        if s.startswith("教材章节："):
            continue
        if s.startswith("## "):
            result["id"] = s[3:].strip()
            continue
        if s.startswith("题型：") or s.startswith("English:"):
            continue
        elif s.startswith("题干：") or s.startswith("题目："):
            result["stem"] = s[3:].strip()
        elif s == "选项：":
            in_options = True
        elif in_options and s.startswith("- "):
            if opt_buf:
                result["options"].append(opt_buf)
            opt_buf = s[2:].strip()
        elif in_options and s.startswith("English:") and opt_buf:
            continue
        elif s.startswith("答案："):
            in_options = False
            if opt_buf:
                result["options"].append(opt_buf)
                opt_buf = None
            result["answer"] = s[3:].strip()
        elif s == "解析：":
            in_analysis = True
            in_options = False
        elif in_analysis and s:
            if s == "---":
                break
            sec_lines.append(s)

    result["analysis_text"] = "\n".join(sec_lines)
    return result if result["stem"] else None


def main():
    parser = argparse.ArgumentParser(description="将润色后的小节 md 转为标准试卷 DOCX")
    parser.add_argument("-i", "--input", default="")
    parser.add_argument("-o", "--output", default="")
    parser.add_argument("--batch", default="", help="批量模式：指定目录，转换所有 p*-ch*-h*.md")
    parser.add_argument(
        "--template",
        default=str(DEFAULT_TEMPLATE_PATH),
        help="题库导入模板DOCX",
    )
    args = parser.parse_args()
    template_path = Path(args.template)

    if args.batch:
        batch_dir = Path(args.batch)
        if not batch_dir.is_dir():
            print(f"错误：目录不存在：{batch_dir}")
            sys.exit(1)
        md_files = sorted(batch_dir.glob("p*-ch*-h*.md"))
        if not md_files:
            print(f"错误：目录下无 p*-ch*-h*.md 文件：{batch_dir}")
            sys.exit(1)
        total_files = len(md_files)
        docx_dir = batch_dir / "docx"
        docx_dir.mkdir(parents=True, exist_ok=True)
        grand_total = 0
        for i, md_path in enumerate(md_files, 1):
            docx_path = docx_dir / md_path.with_suffix(".docx").name
            total, _ = convert_md_to_docx(md_path, docx_path, template_path)
            grand_total += total
            print(f"  [{i}/{total_files}] {md_path.name} → docx/{docx_path.name} ({total}题)")
        print(f"批量完成：{total_files}文件, {grand_total}题, 输出到 {docx_dir}")
        return

    if not args.input or not args.output:
        print("错误：单文件模式需要 -i 和 -o，或使用 --batch 批量模式")
        sys.exit(1)

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        print(f"错误：输入文件不存在：{input_path}")
        sys.exit(1)

    total, skipped = convert_md_to_docx(input_path, output_path, template_path)
    print(f"已保存：{output_path}（共{total}题）")


if __name__ == "__main__":
    main()
